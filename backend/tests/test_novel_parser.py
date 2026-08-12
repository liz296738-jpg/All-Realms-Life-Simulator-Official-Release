"""小说解析：TXT 编码探测 / DOCX / 清洗 / 抽样确定性 / 大小限制。"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest

from game.novel_parser import clean_text, extract_text, sample_text, validate_upload_size


def test_txt_utf8():
    raw = "第一章 开始\n正文内容。".encode("utf-8")
    assert "第一章" in extract_text("a.txt", raw)


def test_txt_gb18030():
    raw = "第一章 开始\n正文内容。".encode("gb18030")
    assert "第一章" in extract_text("b.txt", raw)


def test_txt_utf8_with_bom():
    raw = "﻿第一章 开始".encode("utf-8")
    assert extract_text("c.txt", raw).strip("﻿").startswith("第一章")


def test_docx_extract():
    import io
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一章 山门")
    doc.add_paragraph("少年推开旧门。")
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_text("novel.docx", buf.getvalue())
    assert "第一章 山门" in text
    assert "少年推开旧门" in text


def test_reject_doc():
    with pytest.raises(ValueError):
        extract_text("old.doc", b"abc")


def test_reject_unknown_ext():
    with pytest.raises(ValueError):
        extract_text("novel.pdf", b"abc")


def test_reject_big_file():
    with pytest.raises(ValueError):
        validate_upload_size(b"x" * (30 * 1024 * 1024 + 1))


def test_clean_collapses_blank_lines():
    assert clean_text("a\n\n\n\nb") == "a\n\nb"


def test_clean_normalizes_crlf():
    assert clean_text("a\r\nb\r\nc") == "a\nb\nc"


def test_sample_short_untouched():
    text = "短文本内容" * 20
    assert sample_text(text) == clean_text(text)


def test_sample_within_max_chars_keeps_all():
    text = "x" * 1000
    assert sample_text(text, max_chars=2000) == text


def test_sample_deterministic_and_bounded():
    text = ("甲" * 20000) + ("乙" * 20000) + ("丙" * 20000)
    s1 = sample_text(text)
    s2 = sample_text(text)
    assert s1 == s2  # 同输入恒同输出
    assert len(s1) <= 35000
    assert s1.startswith("甲") and s1.endswith("丙")


def test_sample_long_takes_head_tail_and_middle():
    text = "头部标记" + "内" * 30000 + "尾部标记"
    s = sample_text(text, head=500, tail=300, mid=300, max_chars=1500)
    assert len(s) <= 1500
    assert "头部标记" in s and "尾部标记" in s  # 首尾都在
    assert "……" in s  # 有省略分段标记
