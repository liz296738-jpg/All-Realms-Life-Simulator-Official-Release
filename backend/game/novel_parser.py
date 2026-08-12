"""小说文件解析：TXT（编码探测）+ DOCX 抽文本 + 清洗 + 抽样精读。

抽样精读（确定性）：同一文件永远抽到同一份样本（种子 = sha256 内容），
便于测试与复现；总长 ≤ max_chars 直接用，否则取 开头 + 中间均匀几段 + 结尾。
"""
from __future__ import annotations

import hashlib
import re

_MAX_SIZE = 30 * 1024 * 1024  # 30MB


def extract_text(filename: str, raw: bytes) -> str:
    """按扩展名抽取纯文本。TXT 做编码探测；DOCX 用 python-docx。.doc 抛 ValueError。"""
    name = (filename or "").lower()
    if name.endswith(".docx"):
        return _extract_docx(raw)
    if name.endswith(".doc"):
        raise ValueError("暂不支持旧版 .doc 文档。请在 Word 里另存为 .docx 后再上传。")
    if name.endswith((".txt", ".md")):
        return _decode_txt(raw)
    raise ValueError("仅支持 TXT 或 Word(.docx) 文件")


def _decode_txt(raw: bytes) -> str:
    """按 utf-8-sig → gb18030 → gbk → big5 → latin-1 顺序探测解码。"""
    for enc in ("utf-8-sig", "gb18030", "gbk", "big5", "latin-1"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="ignore")


def _extract_docx(raw: bytes) -> str:
    import io
    from docx import Document  # 延迟导入：无文档需求的环境不拉重依赖

    doc = Document(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def clean_text(text: str) -> str:
    """去 BOM/控制字符、归一化换行、压缩连续空行。"""
    if text.startswith("﻿"):
        text = text[1:]
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def sample_text(text: str, head: int = 15000, tail: int = 10000,
                mid: int = 10000, max_chars: int = 35000) -> str:
    """抽样精读：≤max_chars 直接用；否则 开头 + 均匀中间段 + 结尾。

    抽样种子 = sha256(内容)，同一输入恒得同一输出（确定性、可测）。
    """
    text = clean_text(text)
    if len(text) <= max_chars:
        return text[:max_chars]

    n_mid = 3
    each = max(1, mid // n_mid)
    pool = len(text) - head - tail
    if pool < each * n_mid:
        # 中间不够分：退回 头部+结尾
        return text[:head] + "\n……\n" + text[-tail:]

    seed = int(hashlib.sha256(text.encode("utf-8")).hexdigest(), 16)
    span = pool - each  # 中间窗口可移动范围
    step = max(1, span // n_mid)
    chunks = []
    for i in range(n_mid):
        start = head + i * step + (seed % (step + 1))
        seed //= 7
        start = min(head + span, max(head, start))
        chunks.append(text[start:start + each])

    out = text[:head] + "\n……\n" + "\n……\n".join(chunks) + "\n……\n" + text[-tail:]
    # 分隔符开销可能让拼接略超 max_chars，统一截断到上限
    return out[:max_chars]


def validate_upload_size(raw: bytes) -> None:
    if len(raw) > _MAX_SIZE:
        raise ValueError("文件超过 30MB，请截取部分章节后再上传")
